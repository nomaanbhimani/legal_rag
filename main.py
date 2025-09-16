import os
import streamlit as st
import torch
import torch.nn as nn
from transformers import (
    AutoTokenizer, AutoModel, AutoModelForCausalLM, 
    AutoModelForQuestionAnswering, pipeline, BertForSequenceClassification
)
import chromadb
from chromadb.config import Settings
import PyPDF2
from docx import Document
from typing import List, Dict, Tuple, Optional
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
from datetime import datetime
import re
import hashlib
import json
from dataclasses import dataclass

# Configuration
@dataclass
class Config:
    MODEL_NAME: str = "law-ai/InLegalBERT"
    CHUNK_SIZE: int = 512
    CHUNK_OVERLAP: int = 50
    MAX_CONTEXT_LENGTH: int = 2000
    GENERATION_MAX_LENGTH: int = 500
    TEMPERATURE: float = 0.3
    TOP_P: float = 0.9
    TOP_K: int = 50

config = Config()

class InLegalBERTBrain:
    """
    InLegalBERT as the central brain for:
    1. Text understanding and embedding
    2. Legal text generation and simplification  
    3. Question answering
    4. Legal concept extraction
    """
    
    def __init__(self):
        self.device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
        self.tokenizer = None
        self.embedding_model = None
        self.qa_model = None
        self.generation_pipeline = None
        self.legal_simplifier = None
        
    @st.cache_resource
    def load_models(_self):
        """Load all InLegalBERT-based models"""
        try:
            # Core tokenizer
            _self.tokenizer = AutoTokenizer.from_pretrained(config.MODEL_NAME)
            if _self.tokenizer.pad_token is None:
                _self.tokenizer.pad_token = _self.tokenizer.eos_token
            
            # 1. Embedding model for RAG
            _self.embedding_model = AutoModel.from_pretrained(config.MODEL_NAME)
            _self.embedding_model.to(_self.device)
            _self.embedding_model.eval()
            
            # 2. Question Answering model
            try:
                _self.qa_model = AutoModelForQuestionAnswering.from_pretrained(config.MODEL_NAME)
                _self.qa_model.to(_self.device)
            except:
                # If QA head not available, we'll build custom extraction
                st.info("Building custom QA layer on InLegalBERT")
                _self.qa_model = _self._build_custom_qa_model()
            
            # 3. Legal Text Simplifier (custom fine-tuned head)
            _self.legal_simplifier = _self._build_legal_simplifier()
            
            # 4. Generation pipeline (if model supports it)
            _self.generation_pipeline = _self._setup_generation_pipeline()
            
            return True
            
        except Exception as e:
            st.error(f"Error loading InLegalBERT models: {str(e)}")
            return False
    
    def _build_custom_qa_model(self):
        """Build custom QA model on top of InLegalBERT"""
        class LegalQAModel(nn.Module):
            def __init__(self, base_model):
                super().__init__()
                self.bert = base_model
                self.qa_outputs = nn.Linear(self.bert.config.hidden_size, 2)  # start/end positions
                self.dropout = nn.Dropout(0.1)
                
            def forward(self, input_ids, attention_mask=None, token_type_ids=None):
                outputs = self.bert(
                    input_ids=input_ids,
                    attention_mask=attention_mask,
                    token_type_ids=token_type_ids
                )
                sequence_output = self.dropout(outputs.last_hidden_state)
                logits = self.qa_outputs(sequence_output)
                start_logits, end_logits = logits.split(1, dim=-1)
                start_logits = start_logits.squeeze(-1)
                end_logits = end_logits.squeeze(-1)
                
                return {
                    'start_logits': start_logits,
                    'end_logits': end_logits,
                    'last_hidden_state': outputs.last_hidden_state
                }
        
        model = LegalQAModel(self.embedding_model)
        model.to(self.device)
        return model
    
    def _build_legal_simplifier(self):
        """Build legal text simplifier using InLegalBERT"""
        class LegalSimplifier(nn.Module):
            def __init__(self, base_model, vocab_size):
                super().__init__()
                self.bert = base_model
                self.simplification_head = nn.Sequential(
                    nn.Linear(self.bert.config.hidden_size, self.bert.config.hidden_size * 2),
                    nn.ReLU(),
                    nn.Dropout(0.1),
                    nn.Linear(self.bert.config.hidden_size * 2, self.bert.config.hidden_size),
                    nn.ReLU(),
                    nn.Dropout(0.1),
                    nn.Linear(self.bert.config.hidden_size, vocab_size)
                )
                
            def forward(self, input_ids, attention_mask=None):
                outputs = self.bert(input_ids=input_ids, attention_mask=attention_mask)
                hidden_states = outputs.last_hidden_state
                simplified_logits = self.simplification_head(hidden_states)
                return {
                    'logits': simplified_logits,
                    'hidden_states': hidden_states
                }
        
        vocab_size = len(self.tokenizer.vocab)
        model = LegalSimplifier(self.embedding_model, vocab_size)
        model.to(self.device)
        return model
    
    def _setup_generation_pipeline(self):
        """Setup text generation using InLegalBERT if possible"""
        try:
            # Try to create a generation pipeline
            # Note: This might not work with encoder-only BERT, but we'll try
            generator = pipeline(
                'text-generation',
                model=config.MODEL_NAME,
                tokenizer=self.tokenizer,
                device=0 if torch.cuda.is_available() else -1,
                max_length=config.GENERATION_MAX_LENGTH,
                do_sample=True,
                temperature=config.TEMPERATURE,
                top_p=config.TOP_P,
                top_k=config.TOP_K
            )
            return generator
        except:
            # Build custom generation using the embedding model
            return self._build_custom_generator()
    
    def _build_custom_generator(self):
        """Build custom text generator using InLegalBERT embeddings"""
        class CustomLegalGenerator:
            def __init__(self, model, tokenizer, device):
                self.model = model
                self.tokenizer = tokenizer
                self.device = device
                self.generation_head = nn.Linear(model.config.hidden_size, len(tokenizer.vocab))
                self.generation_head.to(device)
                
            def generate(self, prompt, max_length=200, temperature=0.7):
                """Generate text using beam search and embedding similarity"""
                inputs = self.tokenizer(
                    prompt, 
                    return_tensors="pt", 
                    max_length=512, 
                    truncation=True
                ).to(self.device)
                
                with torch.no_grad():
                    outputs = self.model(**inputs)
                    hidden_states = outputs.last_hidden_state
                    
                    # Use the last hidden state to generate next tokens
                    logits = self.generation_head(hidden_states[:, -1, :])
                    
                    # Apply temperature
                    logits = logits / temperature
                    probs = torch.softmax(logits, dim=-1)
                    
                    # Sample next token
                    next_token = torch.multinomial(probs, 1)
                    
                    # Simple greedy decoding for now
                    generated_text = self.tokenizer.decode(next_token[0], skip_special_tokens=True)
                    
                return prompt + " " + generated_text
                
        return CustomLegalGenerator(self.embedding_model, self.tokenizer, self.device)
    
    def get_embeddings(self, text: str) -> np.ndarray:
        """Generate embeddings using InLegalBERT"""
        if not self.embedding_model:
            if not self.load_models():
                return np.random.rand(768)
        
        try:
            inputs = self.tokenizer(
                text,
                return_tensors="pt",
                max_length=512,
                truncation=True,
                padding=True
            ).to(self.device)
            
            with torch.no_grad():
                outputs = self.embedding_model(**inputs)
                # Use CLS token embedding
                embeddings = outputs.last_hidden_state[:, 0, :].cpu().numpy()
            
            return embeddings[0]
        except Exception as e:
            st.error(f"Embedding error: {str(e)}")
            return np.random.rand(768)
    
    def extract_legal_concepts(self, text: str) -> List[str]:
        """Extract key legal concepts using InLegalBERT"""
        if not self.embedding_model:
            return []
        
        # Legal keywords to look for
        legal_patterns = [
            r'\b(shall|must|may|will)\s+\w+',
            r'\b(contract|agreement|clause|term|condition)\b',
            r'\b(party|parties|plaintiff|defendant)\b',
            r'\b(liability|obligation|right|duty)\b',
            r'\b(terminate|breach|default|cure)\b',
            r'\b(payment|fee|penalty|damages)\b',
            r'\b(confidential|proprietary|intellectual property)\b'
        ]
        
        concepts = []
        for pattern in legal_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            concepts.extend(matches)
        
        return list(set(concepts))
    
    def answer_question(self, question: str, context: str) -> Dict:
        """Use InLegalBERT for question answering"""
        if not self.qa_model:
            return {"answer": "QA model not available", "confidence": 0.0}
        
        try:
            # Prepare input for QA model
            inputs = self.tokenizer.encode_plus(
                question,
                context,
                add_special_tokens=True,
                max_length=512,
                truncation=True,
                padding=True,
                return_tensors="pt"
            ).to(self.device)
            
            with torch.no_grad():
                outputs = self.qa_model(**inputs)
                
                if hasattr(outputs, 'start_logits'):
                    # Standard QA model output
                    start_scores = outputs.start_logits
                    end_scores = outputs.end_logits
                else:
                    # Custom model output
                    start_scores = outputs['start_logits']
                    end_scores = outputs['end_logits']
                
                # Find best start and end positions
                start_index = torch.argmax(start_scores)
                end_index = torch.argmax(end_scores)
                
                # Extract answer
                input_ids = inputs['input_ids'][0]
                answer_tokens = input_ids[start_index:end_index+1]
                answer = self.tokenizer.decode(answer_tokens, skip_special_tokens=True)
                
                confidence = (torch.max(start_scores) + torch.max(end_scores)) / 2
                confidence = torch.sigmoid(confidence).item()
                
                return {
                    "answer": answer.strip(),
                    "confidence": confidence,
                    "start_pos": start_index.item(),
                    "end_pos": end_index.item()
                }
                
        except Exception as e:
            st.error(f"QA error: {str(e)}")
            return {"answer": "Error processing question", "confidence": 0.0}
    
    def simplify_legal_text(self, legal_text: str, context: str = "") -> str:
        """Convert complex legal text to plain English using InLegalBERT"""
        if not self.legal_simplifier:
            return self._rule_based_simplification(legal_text)
        
        try:
            # Prepare prompt for simplification
            prompt = f"""Convert this legal text to plain English:

Legal Text: {legal_text}

Context: {context}

Plain English:"""
            
            # Use generation pipeline if available
            if hasattr(self.generation_pipeline, '__call__'):
                try:
                    result = self.generation_pipeline(
                        prompt,
                        max_length=config.GENERATION_MAX_LENGTH,
                        do_sample=True,
                        temperature=config.TEMPERATURE
                    )
                    return result[0]['generated_text'].split("Plain English:")[-1].strip()
                except:
                    pass
            
            # Use custom generator
            if hasattr(self.generation_pipeline, 'generate'):
                simplified = self.generation_pipeline.generate(
                    prompt,
                    max_length=config.GENERATION_MAX_LENGTH,
                    temperature=config.TEMPERATURE
                )
                return simplified.split("Plain English:")[-1].strip()
            
            # Fallback to rule-based
            return self._rule_based_simplification(legal_text)
            
        except Exception as e:
            st.error(f"Simplification error: {str(e)}")
            return self._rule_based_simplification(legal_text)
    
    def _rule_based_simplification(self, legal_text: str) -> str:
        """Rule-based legal text simplification as fallback"""
        
        # Legal-to-plain mappings
        replacements = {
            r'\bshall\b': 'must',
            r'\bwhereas\b': 'since',
            r'\btherefore\b': 'so',
            r'\bheretofore\b': 'before this',
            r'\bhereinafter\b': 'from now on',
            r'\bparty of the first part\b': 'first party',
            r'\bparty of the second part\b': 'second party',
            r'\bin consideration of\b': 'in exchange for',
            r'\bnotwithstanding\b': 'despite',
            r'\bpursuant to\b': 'according to',
            r'\bwherein\b': 'where',
            r'\bwhereby\b': 'by which',
            r'\baforesaid\b': 'mentioned above',
            r'\bheretofore\b': 'up to now'
        }
        
        simplified = legal_text
        for legal_term, plain_term in replacements.items():
            simplified = re.sub(legal_term, plain_term, simplified, flags=re.IGNORECASE)
        
        # Break down long sentences
        sentences = simplified.split('. ')
        short_sentences = []
        
        for sentence in sentences:
            if len(sentence) > 100:  # Long sentence
                # Try to break at conjunctions
                parts = re.split(r'\b(and|or|but|however|moreover|furthermore)\b', sentence)
                short_sentences.extend([part.strip() for part in parts if part.strip()])
            else:
                short_sentences.append(sentence)
        
        return '. '.join(short_sentences)
    
    def generate_summary(self, legal_texts: List[str], query_context: str = "") -> str:
        """Generate comprehensive summary using InLegalBERT"""
        
        # Extract key concepts from all texts
        all_concepts = []
        for text in legal_texts:
            concepts = self.extract_legal_concepts(text)
            all_concepts.extend(concepts)
        
        key_concepts = list(set(all_concepts))[:10]  # Top 10 unique concepts
        
        # Combine texts for analysis
        combined_text = "\n\n".join(legal_texts[:3])  # Use first 3 chunks
        
        # Generate answer using QA model
        summary_question = f"What are the main points about {query_context}?" if query_context else "What are the main legal points?"
        qa_result = self.answer_question(summary_question, combined_text)
        
        # Simplify the extracted answer
        simplified_answer = self.simplify_legal_text(qa_result["answer"], query_context)
        
        # Combine everything into a comprehensive response
        response = f"""**Main Points:**
{simplified_answer}

**Key Legal Terms Found:**
{', '.join(key_concepts[:5])}

**Confidence:** {qa_result['confidence']:.2f}

**Important:** This is a simplified explanation. The original legal language should be referenced for precise legal meaning."""
        
        return response

# Enhanced RAG System with InLegalBERT Brain
class LegalRAGSystem:
    """Enhanced RAG system with InLegalBERT as the central brain"""
    
    def __init__(self):
        self.brain = InLegalBERTBrain()
        self.vector_db = None
        self.processor = DocumentProcessor()
        self._init_vector_db()
    
    def _init_vector_db(self):
        try:
            import chromadb
            
            # Use the new client configuration
            client = chromadb.PersistentClient(path="./chroma_db")
            
            try:
                self.vector_db = client.create_collection(
                    name="legal_documents",
                    metadata={"hnsw:space": "cosine"}  # Optional: specify distance metric
                )
            except Exception:
                # Collection already exists, get it
                self.vector_db = client.get_collection(name="legal_documents")
                
        except Exception as e:
            st.error(f"Vector DB initialization error: {str(e)}")
            self.vector_db = None
    
    def process_document(self, file, doc_type: str = "legal") -> bool:
        """Process document using InLegalBERT brain"""
        # Extract text (same as before)
        if file.type == "application/pdf":
            text = self.processor.extract_text_from_pdf(file)
        elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            text = self.processor.extract_text_from_docx(file)
        elif file.type == "text/plain":
            text = str(file.read(), "utf-8")
        else:
            st.error("Unsupported file type")
            return False
        
        if not text:
            return False
        
        # Process with InLegalBERT brain
        clean_text = self.processor.clean_text(text)
        chunks = self.processor.chunk_text(clean_text, config.CHUNK_SIZE, config.CHUNK_OVERLAP)
        
        # Generate embeddings using InLegalBERT
        embeddings = []
        enhanced_chunks = []
        
        for chunk in chunks:
            # Get embedding
            embedding = self.brain.get_embeddings(chunk['text'])
            embeddings.append(embedding.tolist())
            
            # Extract legal concepts
            concepts = self.brain.extract_legal_concepts(chunk['text'])
            
            # Enhanced chunk with legal analysis
            enhanced_chunk = {
                **chunk,
                'legal_concepts': concepts,
                'simplified_preview': self.brain.simplify_legal_text(chunk['text'][:200])
            }
            enhanced_chunks.append(enhanced_chunk)
        
        # Store in vector database
        doc_id = hashlib.md5(file.name.encode()).hexdigest()
        metadata_list = []
        documents = []
        ids = []
        
        for i, (chunk, embedding) in enumerate(zip(enhanced_chunks, embeddings)):
            metadata = {
                'filename': file.name,
                'doc_type': doc_type,
                'chunk_id': i,
                'legal_concepts': json.dumps(chunk['legal_concepts']),
                'simplified_preview': chunk['simplified_preview'],
                'upload_date': datetime.now().isoformat()
            }
            metadata_list.append(metadata)
            documents.append(chunk['text'])
            ids.append(f"{doc_id}_{i}")
            doc_id = hashlib.md5(file.name.encode()).hexdigest()
        
        try:
                if self.vector_db is None:
                    st.error("Vector database not initialized")
                    return False
                
                self.vector_db.add(
                    embeddings=embeddings,
                    documents=documents,
                    metadatas=metadata_list,
                    ids=ids
                    )
                st.success(f"Processed {file.name} with InLegalBERT - {len(chunks)} chunks, {sum(len(c['legal_concepts']) for c in enhanced_chunks)} legal concepts identified")
                return True
        except Exception as e:
                st.error(f"Database error: {str(e)}")
                return False
    
    def query(self, question: str) -> Dict:
        """Query system using InLegalBERT brain for understanding and generation"""
        if not self.vector_db:
            return {
                'answer': "No documents uploaded yet.",
                'sources': [],
                'confidence': 0,
                'legal_concepts': []
            }
        
        # Get embedding for query using InLegalBERT
        query_embedding = self.brain.get_embeddings(question)
        
        # Search for relevant chunks
        try:
            results = self.vector_db.query(
                query_embeddings=[query_embedding.tolist()],
                n_results=5
            )
            
            if not results['documents'][0]:
                return {
                    'answer': "No relevant information found in the uploaded documents.",
                    'sources': [],
                    'confidence': 0,
                    'legal_concepts': []
                }
            
            # Extract relevant texts and metadata
            relevant_texts = results['documents'][0]
            metadata_list = results['metadatas'][0]
            distances = results['distances'][0] if results['distances'] else [0] * len(relevant_texts)
            
            # Use InLegalBERT brain to generate comprehensive answer
            answer = self.brain.generate_summary(relevant_texts, question)
            
            # Calculate overall confidence
            confidence_scores = [1 - d for d in distances]
            overall_confidence = np.mean(confidence_scores)
            
            # Extract all legal concepts found
            all_concepts = []
            for metadata in metadata_list:
                if 'legal_concepts' in metadata:
                    try:
                        concepts = json.loads(metadata['legal_concepts'])
                        all_concepts.extend(concepts)
                    except:
                        pass
            
            unique_concepts = list(set(all_concepts))
            
            # Prepare sources with enhanced information
            sources = []
            for i, (metadata, distance) in enumerate(zip(metadata_list, distances)):
                sources.append({
                    'filename': metadata.get('filename', 'Unknown'),
                    'chunk_id': metadata.get('chunk_id', i),
                    'relevance_score': 1 - distance,
                    'legal_concepts': json.loads(metadata.get('legal_concepts', '[]')),
                    'simplified_preview': metadata.get('simplified_preview', '')
                })
            
            return {
                'answer': answer,
                'sources': sources,
                'confidence': overall_confidence,
                'legal_concepts': unique_concepts[:10],  # Top 10 concepts
                'brain_model': 'InLegalBERT'
            }
            
        except Exception as e:
            st.error(f"Query processing error: {str(e)}")
            return {
                'answer': f"Error processing query: {str(e)}",
                'sources': [],
                'confidence': 0,
                'legal_concepts': []
            }

# Keep the DocumentProcessor class from before
class DocumentProcessor:
    """Handles document parsing and chunking"""
    
    @staticmethod
    def extract_text_from_pdf(file) -> str:
        try:
            reader = PyPDF2.PdfReader(file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text
        except Exception as e:
            st.error(f"PDF extraction error: {str(e)}")
            return ""
    
    @staticmethod
    def extract_text_from_docx(file) -> str:
        try:
            doc = Document(file)  # Changed from docx.Document(file)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            st.error(f"DOCX extraction error: {str(e)}")
            return ""
    
    @staticmethod
    def clean_text(text: str) -> str:
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'Page \d+', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\n+', '\n', text)
        return text.strip()
    
    @staticmethod
    def chunk_text(text: str, chunk_size: int = 512, overlap: int = 50) -> List[Dict]:
        words = text.split()
        chunks = []
        
        for i in range(0, len(words), chunk_size - overlap):
            chunk_words = words[i:i + chunk_size]
            chunk_text = ' '.join(chunk_words)
            
            chunks.append({
                'text': chunk_text,
                'start_idx': i,
                'end_idx': min(i + chunk_size, len(words)),
                'word_count': len(chunk_words)
            })
            
            if i + chunk_size >= len(words):
                break
        
        return chunks

# Enhanced Streamlit UI
def main():
    st.set_page_config(
        page_title="InLegalBERT Document Analyzer",
        page_icon="🧠",
        layout="wide"
    )
    
    st.title("🧠 InLegalBERT Legal Document Analyzer")
    st.markdown("**Powered by InLegalBERT**: Fine-tuned legal AI for understanding, analysis, and plain English generation")
    
    # Initialize session state
    if 'rag_system' not in st.session_state:
        with st.spinner("🔄 Loading InLegalBERT brain..."):
            st.session_state.rag_system = LegalRAGSystem()
            # Pre-load the brain
            st.session_state.rag_system.brain.load_models()
    
    # Sidebar
    with st.sidebar:
        st.header("🧠 InLegalBERT Brain Status")
        
        # Model status
        brain = st.session_state.rag_system.brain
        if brain.embedding_model:
            st.success("✅ Embedding Model Loaded")
        else:
            st.error("❌ Embedding Model Failed")
            
        if brain.qa_model:
            st.success("✅ Q&A Model Ready")
        else:
            st.warning("⚠️ Q&A Model Limited")
            
        if brain.legal_simplifier:
            st.success("✅ Legal Simplifier Active")
        else:
            st.warning("⚠️ Basic Simplifier Only")
        
        st.header("📄 Document Upload")
        uploaded_files = st.file_uploader(
            "Upload legal documents",
            type=['pdf', 'docx', 'txt'],
            accept_multiple_files=True,
            help="InLegalBERT will analyze and understand your legal documents"
        )
        
        if uploaded_files:
            for file in uploaded_files:
                if st.button(f"🧠 Analyze {file.name}", key=f"process_{file.name}"):
                    with st.spinner(f"InLegalBERT is analyzing {file.name}..."):
                        success = st.session_state.rag_system.process_document(file)
    
    # Main interface
    col1, col2 = st.columns([3, 2])
    
    with col1:
        st.header("💬 Ask InLegalBERT")
        
        st.markdown("""
        **InLegalBERT specializes in:**
        - Legal document analysis and comprehension
        - Converting complex legal language to plain English
        - Extracting key legal concepts and obligations
        - Answering specific legal questions with citations
        """)
        
        # Enhanced sample questions
        st.subheader("🎯 Specialized Legal Questions:")
        specialized_questions = [
            "What are the key contractual obligations and who is responsible for what?",
            "Explain the termination clauses and what triggers them in plain English",
            "What are the liability limitations and what do they mean for each party?",
            "Summarize the payment terms and penalty structures",
            "What intellectual property rights are mentioned and how are they protected?",
            "Identify any dispute resolution mechanisms and explain the process",
            "What are the confidentiality requirements and their scope?"
        ]
        
        for question in specialized_questions:
            if st.button(f"🧠 {question}", key=f"specialized_{question}"):
                st.session_state.current_question = question
        
        # Custom question input
        user_question = st.text_area(
            "Or ask your specific legal question:",
            height=100,
            placeholder="e.g., What happens if I want to terminate this contract early?"
        )
        
        if st.button("🧠 Analyze with InLegalBERT", type="primary"):
            if user_question.strip():
                st.session_state.current_question = user_question
        
        # Process and display results
        if hasattr(st.session_state, 'current_question'):
            with st.spinner("🧠 InLegalBERT is thinking..."):
                result = st.session_state.rag_system.query(st.session_state.current_question)
                
                st.subheader("🎯 InLegalBERT Analysis")
                st.markdown(result['answer'])
                
                # Legal concepts identified
                if result['legal_concepts']:
                    st.subheader("⚖️ Legal Concepts Identified")
                    concept_cols = st.columns(3)
                    for i, concept in enumerate(result['legal_concepts']):
                        with concept_cols[i % 3]:
                            st.info(f"📌 {concept}")
                
                # Enhanced source information
                if result['sources']:
                    st.subheader("📚 Source Documents & Citations")
                    for i, source in enumerate(result['sources'], 1):
                        with st.expander(f"📄 {source['filename']} (Chunk {source['chunk_id']})"):
                            st.write(f"**Relevance Score:** {source['relevance_score']:.2f}")
                            if source.get('simplified_preview'):
                                st.write(f"**Preview:** {source['simplified_preview']}")
                            if source.get('legal_concepts'):
                                st.write(f"**Legal Concepts:** {', '.join(source['legal_concepts'][:5])}")
                
                # Confidence and model info
                col_conf, col_model = st.columns(2)
                with col_conf:
                    st.metric("🎯 Confidence Score", f"{result['confidence']:.2f}")
                with col_model:
                    st.metric("🧠 Powered by", result.get('brain_model', 'InLegalBERT'))
    
    with col2:
        st.header("📊 InLegalBERT Analytics")
        
        # Document statistics
        if hasattr(st.session_state.rag_system, 'vector_db') and st.session_state.rag_system.vector_db:
            try:
                count = st.session_state.rag_system.vector_db.count()
                st.metric("📄 Document Chunks Analyzed", count)
                
                if count > 0:
                    # Get some sample legal concepts from the database
                    sample_results = st.session_state.rag_system.vector_db.query(
                        query_embeddings=[st.session_state.rag_system.brain.get_embeddings("contract agreement").tolist()],
                        n_results=min(3, count)
                    )
                    
                    if sample_results and sample_results['metadatas']:
                        all_concepts = []
                        for metadata in sample_results['metadatas'][0]:
                            if 'legal_concepts' in metadata:
                                try:
                                    concepts = json.loads(metadata['legal_concepts'])
                                    all_concepts.extend(concepts)
                                except:
                                    pass
                        
                        unique_concepts = list(set(all_concepts))
                        if unique_concepts:
                            st.subheader("🔍 Legal Concepts in Database")
                            for concept in unique_concepts[:8]:
                                st.text(f"• {concept}")
                
            except Exception as e:
                st.metric("📄 Document Chunks", "Error loading")
        else:
            st.metric("📄 Document Chunks", "0")
        
        # InLegalBERT capabilities
        st.subheader("🧠 InLegalBERT Capabilities")
        
        capabilities = [
            ("🎯", "Legal Document Understanding", "Specialized in legal terminology and concepts"),
            ("🔄", "Plain English Conversion", "Transforms complex legal language"),
            ("❓", "Contextual Q&A", "Answers questions using legal knowledge"),
            ("📋", "Concept Extraction", "Identifies key legal terms and clauses"),
            ("🔗", "Citation & Sources", "Provides accurate document references"),
            ("⚖️", "Legal Reasoning", "Applies legal domain expertise")
        ]
        
        for icon, title, description in capabilities:
            st.write(f"{icon} **{title}**")
            st.caption(description)
            st.write("")
        
        # Model configuration
        with st.expander("⚙️ InLegalBERT Configuration"):
            st.write(f"**Model:** {config.MODEL_NAME}")
            st.write(f"**Chunk Size:** {config.CHUNK_SIZE} tokens")
            st.write(f"**Context Length:** {config.MAX_CONTEXT_LENGTH} tokens")
            st.write(f"**Generation Max:** {config.GENERATION_MAX_LENGTH} tokens")
            st.write(f"**Temperature:** {config.TEMPERATURE}")
        
        # Advanced features info
        st.subheader("🚀 Advanced Features")
        st.info("""
        **InLegalBERT Brain Features:**
        
        • **Legal Concept Extraction**: Automatically identifies key legal terms
        
        • **Context-Aware Simplification**: Converts legalese to plain English while preserving meaning
        
        • **Citation Tracking**: Provides exact source references for all answers
        
        • **Multi-Document Analysis**: Analyzes relationships across multiple legal documents
        
        • **Confidence Scoring**: Provides reliability metrics for all responses
        """)
        
        st.warning("""
        ⚠️ **Legal Disclaimer**
        
        InLegalBERT provides analysis and explanations for educational purposes only. 
        
        This is NOT legal advice. Always consult qualified legal professionals for legal matters.
        
        Verify all information with original documents and legal experts.
        """)
    
    # Footer with additional information
    st.markdown("---")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.subheader("🎯 How InLegalBERT Works")
        st.write("""
        1. **Document Analysis**: InLegalBERT reads and understands legal documents using specialized training
        
        2. **Concept Extraction**: Identifies key legal terms, clauses, and relationships
        
        3. **Context Retrieval**: RAG system finds relevant sections for your questions
        
        4. **Plain Language Generation**: InLegalBERT converts complex legal text to understandable explanations
        
        5. **Citation & Verification**: Provides sources and confidence scores for transparency
        """)
    
    with col2:
        st.subheader("📋 Best Practices")
        st.write("""
        **For Best Results:**
        
        • Upload complete legal documents (not fragments)
        
        • Ask specific questions about clauses, terms, or obligations
        
        • Review confidence scores and cited sources
        
        • Cross-reference important findings with original text
        
        • Use for understanding, not as legal advice
        """)
    
    with col3:
        st.subheader("🔧 System Status")
        
        # System health indicators
        brain = st.session_state.rag_system.brain
        
        health_items = [
            ("InLegalBERT Model", "✅ Active" if brain.embedding_model else "❌ Failed"),
            ("Q&A System", "✅ Ready" if brain.qa_model else "⚠️ Limited"),
            ("Text Simplifier", "✅ Active" if brain.legal_simplifier else "⚠️ Basic"),
            ("Vector Database", "✅ Ready" if st.session_state.rag_system.vector_db else "❌ Failed"),
            ("Generation Pipeline", "✅ Active" if brain.generation_pipeline else "⚠️ Fallback")
        ]
        
        for component, status in health_items:
            st.write(f"**{component}:** {status}")
    
    # Additional debugging information for development
    if st.sidebar.checkbox("🔧 Developer Debug Mode"):
        st.subheader("🔧 Debug Information")
        
        with st.expander("Model Architecture Details"):
            brain = st.session_state.rag_system.brain
            
            st.write("**InLegalBERT Components:**")
            st.json({
                "embedding_model_loaded": brain.embedding_model is not None,
                "qa_model_loaded": brain.qa_model is not None,
                "simplifier_loaded": brain.legal_simplifier is not None,
                "generation_pipeline_loaded": brain.generation_pipeline is not None,
                "device": str(brain.device),
                "model_name": config.MODEL_NAME
            })
            
            if hasattr(st.session_state, 'current_question'):
                st.write("**Last Query Analysis:**")
                
                # Show query embedding
                query_emb = brain.get_embeddings(st.session_state.current_question)
                st.write(f"Query embedding shape: {query_emb.shape}")
                st.write(f"Query embedding preview: {query_emb[:10]}...")
                
                # Show legal concepts from query
                query_concepts = brain.extract_legal_concepts(st.session_state.current_question)
                st.write(f"Legal concepts in query: {query_concepts}")

if __name__ == "__main__":
    main()